from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from bs4 import BeautifulSoup
import time
import os

def fetch_chat_content(shared_link):
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")

    driver = webdriver.Chrome(options=chrome_options)
    driver.get(shared_link)

    time.sleep(5)  # Allow page to load (adjust if needed)

    messages = []
    try:
        chat_blocks = driver.find_elements(By.CLASS_NAME, "markdown")
        for idx, block in enumerate(chat_blocks):
            html_content = block.get_attribute("innerHTML").strip()
            if html_content:
                author = "user" if idx % 2 == 0 else "assistant"
                messages.append((author, html_content))
    except Exception as e:
        print(f"Error while parsing messages: {e}")

    driver.quit()
    return messages

def parse_html_content(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    elements = []

    for tag in soup.descendants:
        if tag.name:
            if tag.name.startswith('h') and tag.name[1].isdigit():
                elements.append(('heading', tag.get_text(strip=True), int(tag.name[1])))
            elif tag.name in ['ul', 'ol']:
                for li in tag.find_all('li'):
                    elements.append(('list_item', li.get_text(strip=True)))
            elif tag.name == 'p':
                text = tag.get_text(strip=True)
                if text:
                    elements.append(('paragraph', text))
            elif tag.name == 'pre':
                code = tag.get_text()
                elements.append(('code', code))

    return elements

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run("────────────────────────────────────────────────────────────────────")
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    code_paragraph = doc.add_paragraph()
    code_run = code_paragraph.add_run(code_text)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(46, 46, 46)

    code_paragraph._p.get_or_add_pPr().append(
        parse_xml(r'<w:shd {} w:fill="F7F7F8"/>'.format(nsdecls('w')))
    )
    code_paragraph.paragraph_format.line_spacing = 1.0
    code_paragraph.paragraph_format.space_before = Pt(0)
    code_paragraph.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run("────────────────────────────────────────────────────────────────────")
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)

def generate_docx(messages, output_file):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for author, msg in messages:
        if author == "assistant":
            parts = parse_html_content(msg)

            for part in parts:
                if part[0] == 'heading':
                    text, level = part[1], part[2]
                    heading_style = f'Heading {min(level, 3)}'  # Use Heading 1-3
                    para = doc.add_paragraph(text, style=heading_style)
                elif part[0] == 'paragraph':
                    para = doc.add_paragraph(part[1])
                elif part[0] == 'list_item':
                    para = doc.add_paragraph(part[1], style='List Bullet')
                elif part[0] == 'code':
                    add_code_block(doc, part[1])
                    continue  # Code blocks already handle spacing

                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(6)

    doc.save(output_file)
    print(f"✅ Word document saved successfully to {output_file}")

def generate_word_from_link(shared_link, output_path):
    messages = fetch_chat_content(shared_link)
    if not messages:
        raise Exception("No messages found or failed to fetch the conversation.")
    generate_docx(messages, output_path)
