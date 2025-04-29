from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from bs4 import BeautifulSoup
import time
import os

# ----------------- FETCH CHAT --------------------
def fetch_chat_content(shared_link, logger=print):
    logger("Setting up Chrome options...")
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")

    logger("Launching Chrome driver...")
    driver = webdriver.Chrome(options=chrome_options)

    logger(f"Opening shared link: {shared_link}")
    driver.get(shared_link)

    logger("Waiting for page to load...")
    time.sleep(5)

    messages = []
    try:
        logger("Finding chat blocks...")
        chat_blocks = driver.find_elements(By.CLASS_NAME, "markdown")
        logger(f"Found {len(chat_blocks)} blocks, extracting...")
        for block in chat_blocks:
            html_content = block.get_attribute("innerHTML").strip()
            if html_content:
                messages.append(("assistant", html_content))
    except Exception as e:
        logger(f"Error during message extraction: {e}")

    logger("Closing Chrome driver...")
    driver.quit()
    return messages

# ----------------- PARSE CHAT HTML --------------------
def parse_html_content(html_content, logger=print):
    logger("Parsing HTML content...")
    soup = BeautifulSoup(html_content, 'html.parser')
    elements = []

    for tag in soup.descendants:
        if tag.name:
            if tag.name.startswith('h') and tag.name[1].isdigit():
                elements.append(('heading', tag.get_text(strip=True), int(tag.name[1])))
            elif tag.name in ['ul', 'ol']:
                for li in tag.find_all('li'):
                    elements.append(('list_item', li.get_text(strip=True)))
            elif tag.name == 'p':
                text = tag.get_text(strip=True)
                if text:
                    elements.append(('paragraph', text))
            elif tag.name == 'pre':
                code = tag.get_text()
                elements.append(('code', code))
            elif tag.name == 'table':
                table_data = []
                for row in tag.find_all('tr'):
                    cells = row.find_all(['th', 'td'])
                    cell_texts = [cell.get_text(strip=True) for cell in cells]
                    table_data.append(cell_texts)
                elements.append(('table', table_data))
    return elements

# ----------------- BUILD DOCX --------------------
def add_code_block(doc, code_text, logger=print):
    logger("Adding code block...")
    p = doc.add_paragraph()
    run = p.add_run("────────────────────────────────────────────────────────────────────")
    run.font.color.rgb = RGBColor(200, 200, 200)

    code_paragraph = doc.add_paragraph()
    code_run = code_paragraph.add_run(code_text)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(46, 46, 46)

    code_paragraph._p.get_or_add_pPr().append(
        parse_xml(r'<w:shd {} w:fill="F7F7F8"/>'.format(nsdecls('w')))
    )

def add_table(doc, table_data, logger=print):
    if not table_data:
        logger("No table data found.")
        return
    logger("Adding table...")
    rows = len(table_data)
    cols = max(len(row) for row in table_data)

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            table.rows[i].cells[j].text = cell_text

# ----------------- SAVE FILES --------------------
def generate_docx(messages, output_file, logger=print):
    logger("Creating DOCX file...")
    doc = Document()

    # Set narrow margins (0.5 inches on all sides)
    section = doc.sections[0]
    section.top_margin = Pt(36)  # 0.5 inches (1 inch = 72 points, so 0.5 inches = 36 points)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    # Default style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)  # Slightly smaller font to avoid space issues

    for author, msg in messages:
        if author == "assistant":
            parts = parse_html_content(msg, logger)

            for part in parts:
                if part[0] == 'heading':
                    text, level = part[1], part[2]
                    heading_style = f'Heading {min(level, 3)}'
                    doc.add_paragraph(text, style=heading_style)
                elif part[0] == 'paragraph':
                    doc.add_paragraph(part[1])
                elif part[0] == 'list_item':
                    doc.add_paragraph(part[1], style='List Bullet')
                elif part[0] == 'code':
                    add_code_block(doc, part[1], logger)
                elif part[0] == 'table':
                    add_table(doc, part[1], logger)

    logger(f"Saving DOCX to {output_file}...")
    doc.save(output_file)
    logger("DOCX saved.")
